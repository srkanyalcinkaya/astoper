import pandas as pd
import os
from typing import List, Dict, Optional
from docx import Document
import PyPDF2
import io
from fastapi import UploadFile, HTTPException
import logging

logger = logging.getLogger(__name__)

class FileProcessor:
    """Dosya işleme sınıfı"""
    
    ALLOWED_EXTENSIONS = {
        '.xlsx': 'excel',
        '.xls': 'excel', 
        '.csv': 'csv',
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'docx'
    }
    
    @staticmethod
    def is_allowed_file(filename: str) -> bool:
        """Dosya uzantısı kontrolü"""
        if not filename:
            return False
        
        ext = os.path.splitext(filename)[1].lower()
        return ext in FileProcessor.ALLOWED_EXTENSIONS
    
    @staticmethod
    def get_file_type(filename: str) -> str:
        """Dosya tipini belirle"""
        ext = os.path.splitext(filename)[1].lower()
        return FileProcessor.ALLOWED_EXTENSIONS.get(ext, 'unknown')
    
    @staticmethod
    async def process_excel(file: UploadFile) -> List[Dict]:
        """Excel dosyasını işle"""
        try:
            content = await file.read()
            
            if file.filename.endswith('.csv'):
                df = pd.read_csv(io.BytesIO(content))
            else:
                df = pd.read_excel(io.BytesIO(content))
            
            urls = []
            for column in df.columns:
                for value in df[column].dropna():
                    value_str = str(value).strip()
                    if FileProcessor.is_url(value_str):
                        urls.append({
                            'url': value_str,
                            'source': file.filename,
                            'row': df[df[column] == value].index[0] + 1
                        })
            
            logger.info(f"Excel dosyasından {len(urls)} URL bulundu: {file.filename}")
            return urls
            
        except Exception as e:
            logger.error(f"Excel dosyası işleme hatası: {e}")
            raise HTTPException(status_code=400, detail=f"Excel dosyası işlenemedi: {str(e)}")
    
    @staticmethod
    async def process_pdf(file: UploadFile) -> List[Dict]:
        """PDF dosyasını işle"""
        try:
            content = await file.read()
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            urls = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                found_urls = FileProcessor.extract_urls_from_text(text)
                
                for url in found_urls:
                    urls.append({
                        'url': url,
                        'source': file.filename,
                        'page': page_num + 1
                    })
            
            logger.info(f"PDF dosyasından {len(urls)} URL bulundu: {file.filename}")
            return urls
            
        except Exception as e:
            logger.error(f"PDF dosyası işleme hatası: {e}")
            raise HTTPException(status_code=400, detail=f"PDF dosyası işlenemedi: {str(e)}")
    
    @staticmethod
    async def process_docx(file: UploadFile) -> List[Dict]:
        """DOCX dosyasını işle"""
        try:
            content = await file.read()
            
            doc = Document(io.BytesIO(content))
            urls = []
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text
                found_urls = FileProcessor.extract_urls_from_text(text)
                
                for url in found_urls:
                    urls.append({
                        'url': url,
                        'source': file.filename,
                        'paragraph': para_num + 1
                    })
            
            for table_num, table in enumerate(doc.tables):
                for row_num, row in enumerate(table.rows):
                    for cell in row.cells:
                        text = cell.text
                        found_urls = FileProcessor.extract_urls_from_text(text)
                        
                        for url in found_urls:
                            urls.append({
                                'url': url,
                                'source': file.filename,
                                'table': table_num + 1,
                                'row': row_num + 1
                            })
            
            logger.info(f"DOCX dosyasından {len(urls)} URL bulundu: {file.filename}")
            return urls
            
        except Exception as e:
            logger.error(f"DOCX dosyası işleme hatası: {e}")
            raise HTTPException(status_code=400, detail=f"DOCX dosyası işlenemedi: {str(e)}")
    
    @staticmethod
    def is_url(text: str) -> bool:
        """Metin URL mi kontrol et"""
        import re
        url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
        return bool(re.match(url_pattern, text))
    
    @staticmethod
    def extract_urls_from_text(text: str) -> List[str]:
        """Metinden URL'leri çıkar"""
        import re
        url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
        urls = re.findall(url_pattern, text)
        return list(set(urls))  # Tekrarları kaldır
    
    @staticmethod
    def extract_emails_from_text(text: str) -> List[str]:
        """Metinden email adreslerini çıkar"""
        import re
        
        if not text or not text.strip():
            return []
        
        # Word boundary olmadan email pattern - daha kapsamlı
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        
        # Metindeki tüm email adreslerini bul
        emails = re.findall(email_pattern, text)
        
        valid_emails = []
        for email in emails:
            email = email.strip().lower()
            if FileProcessor.is_valid_email(email):
                valid_emails.append(email)
        
        return list(set(valid_emails))  # Tekrarları kaldır
    
    @staticmethod
    def is_valid_email(email: str) -> bool:
        """Email adresi geçerli mi kontrol et"""
        if not email or len(email) < 5:
            return False
        
        file_extensions = [
            '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg', '.webp', '.ico',
            '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx',
            '.zip', '.rar', '.tar', '.gz', '.mp4', '.avi', '.mov', '.wmv',
            '.mp3', '.wav', '.flac', '.css', '.js', '.html', '.htm', '.xml',
            '.json', '.txt', '.log', '.sql', '.php', '.asp', '.jsp'
        ]
        
        email_lower = email.lower()
        for ext in file_extensions:
            if ext in email_lower:
                return False
        
        invalid_patterns = [
            'localhost', '127.0.0.1',
            'www.', 'http://', 'https://', 'ftp://',
            'data:', 'javascript:', 'mailto:', 'tel:'
        ]
        
        for pattern in invalid_patterns:
            if pattern in email_lower:
                return False
        
        import re
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(email_pattern, email):
            return False
        
        if len(email) > 100:
            return False
        
        return True
    
    @staticmethod
    def process_file_from_path(file_path: str) -> Dict:
        """Dosya yolundan email adreslerini çıkar - sync versiyon"""
        try:
            if not os.path.exists(file_path):
                raise Exception(f"Dosya bulunamadı: {file_path}")
            
            filename = os.path.basename(file_path)
            if not FileProcessor.is_allowed_file(filename):
                raise Exception("Desteklenmeyen dosya formatı")
            
            file_type = FileProcessor.get_file_type(filename)
            emails = []
            
            if file_type == 'excel' or file_type == 'csv':
                try:
                    logger.info(f"Excel/CSV dosyası işleniyor: {file_path}")
                    
                    if file_path.endswith('.csv'):
                        df = pd.read_csv(file_path, encoding='utf-8')
                    else:
                        # Excel için farklı engine'ler dene
                        try:
                            df = pd.read_excel(file_path, engine='openpyxl')
                        except Exception as e1:
                            logger.warning(f"OpenPyXL ile okunamadı: {e1}, xlrd ile deneniyor...")
                            try:
                                df = pd.read_excel(file_path, engine='xlrd')
                            except Exception as e2:
                                logger.error(f"XLRD ile de okunamadı: {e2}")
                                raise e1
                    
                    logger.info(f"Excel/CSV okundu: {len(df)} satır, {len(df.columns)} sütun")
                    
                    for column in df.columns:
                        for value in df[column].dropna():
                            value_str = str(value).strip()
                            if value_str:  # Boş değerleri atla
                                found_emails = FileProcessor.extract_emails_from_text(value_str)
                                emails.extend(found_emails)
                                if found_emails:
                                    logger.debug(f"Sütun '{column}' değer '{value_str}' den {len(found_emails)} email bulundu")
                    
                    logger.info(f"Excel/CSV'den toplam {len(emails)} email adresi çıkarıldı")
                    
                except Exception as e:
                    logger.error(f"Excel/CSV okuma hatası: {e}")
                    logger.error(f"Dosya yolu: {file_path}")
                    # Hata durumunda boş liste döndür
                    emails = []
            
            elif file_type == 'pdf':
                try:
                    logger.info(f"PDF dosyası işleniyor: {file_path}")
                    
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        logger.info(f"PDF okundu: {len(pdf_reader.pages)} sayfa")
                        
                        for page_num, page in enumerate(pdf_reader.pages):
                            text = page.extract_text()
                            if text and text.strip():  # Boş sayfaları atla
                                found_emails = FileProcessor.extract_emails_from_text(text)
                                emails.extend(found_emails)
                                if found_emails:
                                    logger.debug(f"Sayfa {page_num + 1}'den {len(found_emails)} email bulundu")
                            else:
                                logger.debug(f"Sayfa {page_num + 1} boş veya metin çıkarılamadı")
                    
                    logger.info(f"PDF'den toplam {len(emails)} email adresi çıkarıldı")
                    
                except Exception as e:
                    logger.error(f"PDF okuma hatası: {e}")
                    logger.error(f"Dosya yolu: {file_path}")
                    # Hata durumunda boş liste döndür
                    emails = []
            
            elif file_type == 'docx':
                try:
                    doc = Document(file_path)
                    for paragraph in doc.paragraphs:
                        found_emails = FileProcessor.extract_emails_from_text(paragraph.text)
                        emails.extend(found_emails)
                    
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                found_emails = FileProcessor.extract_emails_from_text(cell.text)
                                emails.extend(found_emails)
                except Exception as e:
                    logger.error(f"DOCX okuma hatası: {e}")
            
            unique_emails = list(set(emails))
            unique_emails.sort()
            
            logger.info(f"Dosyadan {len(unique_emails)} benzersiz email adresi bulundu: {filename}")
            
            return {
                'emails': unique_emails,
                'file_path': file_path,
                'file_type': file_type,
                'total_found': len(unique_emails)
            }
            
        except Exception as e:
            logger.error(f"Dosya işleme hatası {file_path}: {e}")
            logger.error(f"Dosya türü: {file_type}, Dosya boyutu: {os.path.getsize(file_path) if os.path.exists(file_path) else 'N/A'}")
            raise Exception(f"Dosya işleme hatası: {str(e)}")
    
    @staticmethod
    async def process_file(file: UploadFile) -> List[Dict]:
        """Dosyayı işle - async versiyon (URL çıkarma için)"""
        if not FileProcessor.is_allowed_file(file.filename):
            raise HTTPException(
                status_code=400, 
                detail="Desteklenmeyen dosya formatı. Sadece Excel, CSV, PDF ve DOCX dosyaları kabul edilir."
            )
        
        file_type = FileProcessor.get_file_type(file.filename)
        
        if file_type == 'excel':
            return await FileProcessor.process_excel(file)
        elif file_type == 'pdf':
            return await FileProcessor.process_pdf(file)
        elif file_type == 'docx':
            return await FileProcessor.process_docx(file)
        else:
            raise HTTPException(status_code=400, detail="Desteklenmeyen dosya formatı")

